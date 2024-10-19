from docx import Document
from docx.text.run import Font 
from docx.oxml.text.run import CT_R
from docx.enum.text import WD_ALIGN_PARAGRAPH
path_sample_FL_and_IP = "sample_docx/ФЛ (ИП)_Принимаем арест_ФССП.docx"
path_sample_UL = "sample_docx/ЮЛ_Принимаем арест_ ФССП.docx"
path_answer_FL_and_IP = "__answer_docx/FL_and_IP_otvet.docx"
path_answer_UL = "__answer_docx/UL_otvet.docx"

# функция удаления абзаца в word документе, чтобы не было пустой строки
def delete_paragraph(paragraph): 
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

# поиск в таблице вордфайла шаблонного номера таска и замена на нужный номер
def obrabotka_task_nomer(doc, nomer):
	t = doc.tables
	tc = t[0].cell(0,0).paragraphs
	for i in range(len(tc)):
		s = tc[i].text
		if s != '':
			g = tc[i].runs
			for j in range(len(g)):
				if g[j].text == '5-25WXG79Y8':
					g[j].text = nomer
					tc[i].alignment = WD_ALIGN_PARAGRAPH.LEFT
			g[j].font.highlight_color = None

# поиск в таблице вордфайла шаблонного имени и адреса УФССП
# и замена этих значений на нужное
def obrabotka_name_and_adress_of_shit(doc, ufssp_name, ufssp_adress):
	t = doc.tables
	tc = t[0].cell(0,1).paragraphs
	for i in range(len(tc)):
		g = tc[i].runs
		for j in range(len(g)):
			if g[j].text == 'УФССП по ':
				g[j].text = ufssp_name
			elif g[j].text == '644053, г. Москва, ул. Пушкина, д. 85, корп. 3':
				g[j].text = ufssp_adress
			elif g[j].text == 'Омской области' or g[j].text == 'МОСП' or g[j].text == ' по ' or g[j].text == 'ОВИП':
				g[j].text = ''
			else:
				pass
			g[j].font.highlight_color = None
	delete_paragraph(tc[1])
	
# обработка номера ИП для определения самого номера и окончания (-ИП или -СП)
def obrabotka_ip_nomer(ind_pred_nomer):
	ind_pred_nomer_post = ''
	if ind_pred_nomer[-3:] == '-ИП' or ind_pred_nomer[-3:] == '-СП':
		ind_pred_nomer_post = ind_pred_nomer[-3:]
		ind_pred_nomer = ind_pred_nomer[:-3]
	return ind_pred_nomer, ind_pred_nomer_post

# поиск в первом абзаце ФЛ вордфайла шаблонных
# даты создания документа, номера ип, типа ип, имени человека, даты рождения, номера счета и даты поступления
# и замены их на нужные значения
def obrabotka_first_paragraph_FL_and_IP(doc, data_create_doc, ind_pred_nomer, ind_pred_nomer_post, person_name, birthday,  a_ls, data_postuplenya, count_for_ls):
	for i in range(len(doc.paragraphs)):
		s = doc.paragraphs[i].text
		if s != '':
			g = doc.paragraphs[i].runs
			for j in range(len(g)):
				print(g[j].text)
				if g[j].text == "27.04.2017":
					g[j].text = data_create_doc
				elif g[j].text == "127802/16/18022":
					g[j].text = ind_pred_nomer
				elif g[j].text == '-ИП' or g[j].text == '-СП':
					if ind_pred_nomer_post != '':
						g[j].text = ind_pred_nomer_post
				elif g[j].text == "ИвановаИванаИвановича":
					g[j].text = person_name
				elif g[j].text == "01011990":
					g[j].text = birthday
				elif g[j].text == ", находящиеся на счёте №":
					if count_for_ls == 1:
						g[j].text = ", находящиеся на счёте №"
					else:
						g[j].text = ", находящиеся на счетах №"
				elif g[j].text == "40702810110000031570":
					g[j].text = a_ls;
				elif g[j].text == '13.07.2019':
					g[j].text = data_postuplenya
				else:
					pass
				g[j].font.highlight_color = None

# поиск в первом абзаце ФЛ вордфайла шаблонных
# даты создания документа, номера ип, типа ип, имени человека, даты рождения, номера счета и даты поступления
# и замены их на нужные значения
def obrabotka_first_paragraph_UL(doc, data_create_doc, ind_pred_nomer, ind_pred_nomer_post, ooo_name, inn, a_ls, data_postuplenya, count_for_ls):
	for i in range(len(doc.paragraphs)):
		s = doc.paragraphs[i].text
		if s != '':
			g = doc.paragraphs[i].runs
			for j in range(len(g)):
				if g[j].text == "27.04.2017":
					g[j].text = data_create_doc
				elif g[j].text == "127802/16/18022":
					g[j].text = ind_pred_nomer
				elif g[j].text == '-ИП' or g[j].text == '-СП':
					if ind_pred_nomer_post != '':
						g[j].text = ind_pred_nomer_post
				elif g[j].text == "ОООРОМАШКА":
					g[j].text = ooo_name
				elif g[j].text == "590609527002":
					g[j].text = inn
				elif g[j].text == ", находящиеся на счёте №":
					if count_for_ls == 1:
						g[j].text = ", находящиеся на счёте №"
					else:
						g[j].text = ", находящиеся на счетах №"
				elif g[j].text == "40702810110000031570":
					g[j].text = a_ls;
				elif g[j].text == '13.07.2019':
					g[j].text = data_postuplenya
				else:
					pass
				g[j].font.highlight_color = None

# поиск в вордфайле 2, 3, 4 абзацов, дублирование их runs'ов и их позиции в память и выдача этих значений 
def take_last3_paragraphs(doc):
	for i in range(len(doc.paragraphs)):
		s = doc.paragraphs[i].text
		if s.find("На основании вышеуказанного Постановления") != -1 and s.find("в сумме") != -1:
			pattern0r = doc.paragraphs[i].runs
			pattern0 = doc.paragraphs[i]
			doc.paragraphs[i].clear()

		if s.find("По состоянию на") != -1:
			pattern1r = doc.paragraphs[i].runs
			pattern1 = doc.paragraphs[i]
			doc.paragraphs[i].clear()
		if s.find("На основании вышеуказанного Постановления") != -1 and s.find("арест.") != -1:
			pattern2r = doc.paragraphs[i].runs
			pattern2 = doc.paragraphs[i]
			doc.paragraphs[i].clear()
	return pattern0, pattern0r, pattern1, pattern1r, pattern2, pattern2r

# разделение номеров лс из описания творка на строки без null
def separation_na_stroki(mnogo_strok):
	s_stroki = []
	stroka = ''
	flag = False
	stop = 0
	for i in range(len(mnogo_strok)-4):
		if mnogo_strok[i:i+4] == 'null':
			flag = True
		if mnogo_strok[i] == '\n':
			if flag == False:
				s_stroki.append(stroka + "  ")
			stroka = ''
			flag = False
			stop = i + 1
		else:
			stroka = stroka + mnogo_strok[i]

	if mnogo_strok[stop:].find('null') == -1:
		s_stroki.append(mnogo_strok[stop:] + "  ")
	return s_stroki

# разделение номеров лс из описания творка на номера лс и суммы
def take_lists_ls_and_summa(mnogo_strok):
	s_stroki = separation_na_stroki(mnogo_strok)
	ls = []
	summa = []
	for j in range(len(s_stroki)):
		s = s_stroki[j]
		stage = 0
		element = ''
		flag_num = False
		for i in range(len(s)):
			if flag_num == False and element != '':
				if stage == 0:
					ls.append(element)
					element = ''
					stage = 1
				elif stage == 1:
					summa.append(element)
					element = ''
					stage = 0

			if s[i].isdigit() == True or s[i] == '.' or s[i] == '-':
				flag_num = True
				element = element + s[i]
			else:
				flag_num = False
	return ls, summa

# разделение номеров лс и их сумм на нулевые(отрицательные) и положительные счета
def separation_na_womls_and_wmls(ls, summa, value_type, unic_value_type):
	womls = []
	wmls = []
	for i in range(len(unic_value_type)):
		womls.append([])
		wmls.append([])

	for i in range(len(ls)):
		for j in range(len(unic_value_type)):
			if float(summa[i]) <= 0.0 and unic_value_type[j] == value_type[i]:
				womls[j].append(ls[i])
			elif float(summa[i]) > 0.0 and unic_value_type[j] == value_type[i]:
				wmls[j].append([ls[i], float(summa[i])])
	return womls, wmls

def add_paragraph_with_wmls(wmls, pattern0, pattern0r, data_postuplenya, unic_value_type):
	for i in range(len(wmls)):
		if wmls[i] != []:
			for j in range(len(wmls[i])):
				r = pattern0r
				fnt_name = r[0].font.name
				g = pattern0.insert_paragraph_before()
				g.add_run(f'На основании вышеуказанного Постановления {data_postuplenya} на счет №{wmls[i][j][0]} наложен арест в сумме {wmls[i][j][1]} {unic_value_type[i]}.')
				g.runs[0].font.name = fnt_name
				g.runs[0].bold = r[0].bold
				g.runs[0].style = r[0].style
				g.runs[0].underline = r[0].underline
				g.runs[0].font.highlight_color = None

def add_paragraph_with_womls_summa0(womls, pattern1, pattern1r, data_postuplenya, unic_value_type):
	for i in range(len(womls)):
		if womls[i] != []:
			alot_ls = ', '.join(womls[i])
			r = pattern1r
			fnt_name = r[0].font.name
			g = pattern1.insert_paragraph_before()
			sc = ''
			if len(womls[i]) == 1:
				sc = 'счете №'
			else:
				sc = 'счетах №'
			g.add_run(f'По состоянию на {data_postuplenya} остаток денежных средств на {sc}{alot_ls} составил    0.00 {unic_value_type[i]}.')
			g.runs[0].font.name = fnt_name
			g.runs[0].bold = r[0].bold
			g.runs[0].style = r[0].style
			g.runs[0].underline = r[0].underline
			g.runs[0].font.highlight_color = None

def add_paragraph_with_womls_arest(womls, pattern2, pattern2r, data_postuplenya):
	count = 0
	for i in range(len(womls)):
		if womls[i] == []:
			count += 1
	if count == len(womls):
		return
	alot_ls = ''

	count = 0
	print(womls)
	for i in range(len(womls)):
		alot_ls = alot_ls + ', '.join(womls[i]) + ', '
		for j in range(len(womls[i])):
			count += 1

	alot_ls = alot_ls[:-2]
	r = pattern2r
	fnt_name = r[0].font.name
	g = pattern2.insert_paragraph_before()
	sc = ''
	if count == 1:
		sc = 'счет №'
	else:
		sc = 'счета №'
	g.add_run(f'На основании вышеуказанного Постановления {data_postuplenya} на {sc}{alot_ls} наложен арест.')
	g.runs[0].font.name = fnt_name
	g.runs[0].bold = r[0].bold
	g.runs[0].style = r[0].style
	g.runs[0].underline = r[0].underline
	g.runs[0].font.highlight_color = None

def obrabotka_stroka_data(stroka):
	stroka = stroka.strip()
	for i in range(len(stroka)):
		if stroka[i] == ',':
			stroka = stroka[:i] + '.' + stroka[i+1:]

	return stroka

def print_result(womls, wmls, unic_value_type):
	for i in range(len(womls)):
		if womls[i] == []:
			print("pusto")
		else:
			print(f" {womls[i]} : {unic_value_type[i]}")

	print("С суммой")
	for i in range(len(wmls)):
		for j in range(len(wmls[i])):
			if wmls[i] == []:
				print("pusto")
			else:
				print(f" {wmls[i][j]} : {unic_value_type[i]}")


for_copy = """Номер счета: 40 На счете: 0
Номер счета: 411 На счете: 2
Номер счета: 4222 На счете: 0
Номер счета: 43333 На счете: 322
Номер счета: 444444 На счете: 0
Номер счета: 4555555 На счете: -100.21
Номер счета: null На счете: null
Номер счета: 40802810200001529859 На счете: 0
Номер счета: null На счете: null"""



def all_together(FL_and_IP_or_UL, list_of_data, ls, summa, value_type):
	for i in range(len(list_of_data)):
		print(list_of_data[i])
	name_of_shit = list_of_data[0].strip()
	adress_of_shit = list_of_data[1].strip()
	data_create_doc = obrabotka_stroka_data(list_of_data[2])
	data_postuplenya = obrabotka_stroka_data(list_of_data[3])
	ind_pred_nomer = list_of_data[4].strip()
	ind_pred_name = list_of_data[5].strip()
	inn = list_of_data[6].strip()
	nomer = list_of_data[7].strip()
	for_copy = list_of_data[8]


	count_for_ls = len(ls)
	a_ls = ', '.join(ls)
		
	unic_value_type = list(set(value_type))


	womls = []
	wmls = []
	womls, wmls = separation_na_womls_and_wmls(ls, summa, value_type, unic_value_type)
	print(f"name of shit = {name_of_shit}")
	print(f"type = {type(name_of_shit)}")

	if FL_and_IP_or_UL == "FL_and_IP":
		document = Document(path_sample_FL_and_IP)
	elif FL_and_IP_or_UL == "UL":
		document = Document(path_sample_UL)
	else:
		return None
	

	pattern0, pattern0r, pattern1, pattern1r, pattern2, pattern2r = take_last3_paragraphs(document)

	obrabotka_task_nomer(document, nomer)
	obrabotka_name_and_adress_of_shit(document, name_of_shit, adress_of_shit)
	ind_pred_nomer, ind_pred_nomer_post = obrabotka_ip_nomer(ind_pred_nomer)
	if FL_and_IP_or_UL == "FL_and_IP":
		obrabotka_first_paragraph_FL_and_IP(document,data_create_doc, ind_pred_nomer, ind_pred_nomer_post, ind_pred_name, inn,  a_ls, data_postuplenya, count_for_ls)
	elif FL_and_IP_or_UL == "UL":
		obrabotka_first_paragraph_UL(document, data_create_doc, ind_pred_nomer, ind_pred_nomer_post, ind_pred_name, inn, a_ls, data_postuplenya, count_for_ls)
	else:
		return None

	add_paragraph_with_wmls(wmls, pattern0, pattern0r, data_postuplenya, unic_value_type)
	add_paragraph_with_womls_summa0(womls, pattern1, pattern1r, data_postuplenya, unic_value_type)
	add_paragraph_with_womls_arest(womls, pattern2, pattern2r, data_postuplenya)

	delete_paragraph(pattern0)
	delete_paragraph(pattern1)
	delete_paragraph(pattern2)

	if FL_and_IP_or_UL == "FL_and_IP":
		document.save(path_answer_FL_and_IP)
	elif FL_and_IP_or_UL == "UL":
		document.save(path_answer_UL)
	else:
		return None
